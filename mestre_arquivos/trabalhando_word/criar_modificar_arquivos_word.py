from docx import Document
from docx.shared import Cm


meu_word = Document()
meu_word.add_heading('Titulo do documento', 0)

paragrafo = meu_word.add_paragraph('Um paragrafo simples e ')
paragrafo.add_run('en negrito').bold = True
paragrafo.add_run(' com palavras em ')
paragrafo.add_run('italico').italic = True

meu_word.add_heading('Titulo nivel 1', level=1)
meu_word.add_heading('Titulo nivel 2', level=2)
meu_word.add_heading('Titulo nivel 3', level=3)

meu_word.add_paragraph('Formatação "No Spacing"',style='No Spacing') 
meu_word.add_paragraph('Formatação "Heading 1"', style='Heading 1') 
meu_word.add_paragraph('Formatação "Heading 2"', style='Heading 2') 
meu_word.add_paragraph('Formatação "Heading 3"', style='Heading 3') 
meu_word.add_paragraph('Formatação "Title"', style='Title') 
meu_word.add_paragraph('Formatação "Subtitle"', style='Subtitle') 
meu_word.add_paragraph('Formatação "Quote"', style='Quote') 
meu_word.add_paragraph('Formatação "Intense Quote"', style='Intense Quote') 
meu_word.add_paragraph('Formatação "List Paragraph"', style='List Paragraph') 
meu_word.add_paragraph('Primeiro item em uma lista com pontos ', style='List Bullet') 
meu_word.add_paragraph('Primeiro item em uma lista numerada ', style='List Number')

meu_word.add_picture('russian_skull.jpg', width=Cm(5.25))

# tabela = meu_word.add_table(rows=3, cols=2)
# celula = tabela.cell(0,0)
# celula.text = 'Nome'

registros = (
    (3, '101', 'Maça'),
    (7, '422', 'Ovos'),
    (4, '631', 'Banana'),
)

meu_word.add_page_break()

tabela = meu_word.add_table(rows=1, cols=3)

cabecalho = tabela.rows[0].cells
cabecalho[0].text = 'Quantidade'
cabecalho[1].text = 'Id'
cabecalho[2].text = 'Descrição'

for quantidade, id, descricao in registros:
    dados_por_linha = tabela.add_row().cells
    dados_por_linha[0].text = str(quantidade)
    dados_por_linha[1].text = str(id)
    dados_por_linha[2].text = str(descricao)

meu_word.save('teste.docx')