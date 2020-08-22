from docx import Document

meu_word = Document('teste.docx')

def mostrando_todos_os_paragrafos_de_um_documento():
    paragrafos = meu_word.paragraphs

    for paragrafo in paragrafos:
        print(paragrafo.text)

def alterando_o_valor_de_um_paragrafo():
    print(meu_word.paragraphs[0].text)
    meu_word.paragraphs[0].text = 'Titulo do Documento'

    meu_word.save('teste.docx')
    print(meu_word.paragraphs[0].text)

def acessando_partes_de_paragrafo():
    for indice, run in enumerate(meu_word.paragraphs[1].runs):
        print(f'{indice} {run.text}')
        
    meu_word.paragraphs[1].runs[1].text = 'EM NEGRITO'
    meu_word.save('teste.docx')


if __name__ == '__main__':
    # mostrando_todos_os_paragrafos_de_um_documento()
    # alterando_o_valor_de_um_paragrafo()
    # acessando_partes_de_paragrafo()

    print('** FIM **')