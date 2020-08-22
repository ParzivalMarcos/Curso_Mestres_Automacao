from docx import Document
from docx.shared import Cm

document = Document()

document.add_heading('Olá Mestres', level=0)

primeiro_paragrafo = document.add_paragraph('Bem vindo a comunicade ')
primeiro_paragrafo.add_run('Mestres da Automação ').bold = True
primeiro_paragrafo.add_run('o ligar onde os mestres se encontram.')

document.add_heading('Módulo 1 - Mestre do Software', level=1)
document.add_heading('Módulo 2 - Mestre do Bootcamp Python', level=2)
document.add_heading('Módulo 3 - Mestre da web', level=3)

segundo_paragrafo = document.add_paragraph(
    '"Descubra como a programação pode mudar o futuro e a sua vida"', 
    style='Quote')

document.add_paragraph('Python', style='List Bullet')
document.add_paragraph('Python3', style='List Bullet')
document.add_paragraph('Pip', style='List Bullet')
document.add_paragraph('Pip3', style='List Bullet')

document.add_paragraph('Windows', style='List Number')
document.add_paragraph('Linux', style='List Number')
document.add_paragraph('MacOS', style='List Number')

valores_tabela = (
    (1, 'Instalando Python', '3:05'),
    (2, 'Instalando VS Code', '5:50'),
    (3, 'Instalando Selenium', '3:50'),
)

tabela = document.add_table(rows=1, cols=3)

titulos_colunas = tabela.rows[0].cells
titulos_colunas[0].text = 'Aula'
titulos_colunas[1].text = 'Nome'
titulos_colunas[2].text = 'Duração'

for aula, nome, duracao in valores_tabela:
    linha = tabela.add_row().cells
    linha[0].text = str(aula)
    linha[1].text = nome
    linha[2].text = duracao

document.add_picture('image.jpg', width=Cm(5.25))


document.save('docDesafio.docx')